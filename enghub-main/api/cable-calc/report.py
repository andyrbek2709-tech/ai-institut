"""
Vercel Python serverless function — генерация Word-отчёта по результату расчёта.

POST /api/cable-calc/report
JSON: {
  "input": { ...исходные данные формы... },
  "result": { ...ответ /api/cable-calc/calc... }
}

Ответ: .docx файл с заголовками "Content-Disposition: attachment".
"""
import io
import json
import os
import time
import traceback
from http.server import BaseHTTPRequestHandler

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    Document = None
    _IMPORT_ERR = str(e)


_MATERIAL = {"Cu": "медь", "Al": "алюминий"}
_INSULATION = {"PVC": "ПВХ (70°C)", "XLPE": "сшитый ПЭ (90°C)"}
_METHOD = {
    "A1": "A1 — изолир. провода в трубе в стене",
    "A2": "A2 — кабель в трубе в стене",
    "B1": "B1 — изолир. провода в трубе по стене",
    "B2": "B2 — кабель в трубе по стене",
    "C":  "C — открыто на стене / лотке",
    "D1": "D1 — кабель в трубе в земле",
    "D2": "D2 — кабель напрямую в земле",
    "E":  "E — кабель в воздухе на лотке",
    "F":  "F — кабель в воздухе на крюках",
    "G":  "G — лесенка/полки в воздухе",
}


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def _add_kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = str(k)
        t.rows[i].cells[1].text = str(v)
    return t


def _safe_num(v, decimals=2, suffix=""):
    if v is None or v == "":
        return "—"
    try:
        if decimals is None:
            return f"{v}{suffix}"
        return f"{float(v):.{decimals}f}{suffix}"
    except Exception:
        return f"{v}{suffix}"


def _build_doc(payload: dict) -> bytes:
    inp = payload.get("input", {}) or {}
    res = payload.get("result", {}) or {}

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Заголовок
    title = doc.add_heading("Расчёт кабельной линии до 1 кВ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Методика: МЭК 60364-5-52, СП РК 3.03-104-2014, ПУЭ. "
        "ТКЗ — метод Беляева Е.Н."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(time.strftime("%d.%m.%Y %H:%M")).italic = True

    if inp.get("line_id") or inp.get("line_name"):
        doc.add_heading(
            f"Линия: {inp.get('line_id', '')} {inp.get('line_name', '')}".strip(),
            level=1,
        )

    # 1. Исходные данные
    _add_heading(doc, "1. Исходные данные", level=1)
    _add_kv_table(doc, [
        ("Режим расчёта", {
            "select": "Подбор сечения",
            "check": "Проверка сечения",
            "max_load": "Максимальная нагрузка",
        }.get(inp.get("mode"), inp.get("mode", "—"))),
        ("Число фаз", inp.get("phases", "—")),
        ("Мощность P, кВт", _safe_num(inp.get("power_kw"), 2)),
        ("cos φ", _safe_num(inp.get("cos_phi"), 2)),
        ("Длина L, м", _safe_num(inp.get("length_m"), 1)),
        ("Допустимое ΔU, %", _safe_num(inp.get("delta_u_pct_max"), 2)),
        ("Материал жилы", _MATERIAL.get(inp.get("material"), inp.get("material", "—"))),
        ("Изоляция", _INSULATION.get(inp.get("insulation"), inp.get("insulation", "—"))),
        ("Способ прокладки", _METHOD.get(inp.get("method"), inp.get("method", "—"))),
        ("Темп. среды, °C", _safe_num(inp.get("ambient_temp_c"), 1)),
        ("Кабелей в группе", inp.get("cables_nearby", 1)),
        ("Параллельных кабелей", inp.get("cable_count", 1)),
        ("U_ном, В", _safe_num(inp.get("u_nom_v"), 0)),
        ("Z тр-ра, мОм", _safe_num(inp.get("z_t_mohm"), 1)),
    ])

    # 2. Результаты
    _add_heading(doc, "2. Результаты расчёта", level=1)
    _add_kv_table(doc, [
        ("Расчётный ток I_расч, А", _safe_num(res.get("i_calc_a"), 2)),
        ("Сечение фаза, мм²", _safe_num(res.get("section_mm2"), 0)),
        ("Сечение нулевого, мм²", _safe_num(res.get("section_zero_mm2"), 0)),
        ("Допустимый ток I_доп, А", _safe_num(res.get("i_allowable_a"), 2)),
        ("Падение напряжения ΔU, %", _safe_num(res.get("delta_u_pct"), 3)),
        ("Ток 1ф КЗ, А", _safe_num(res.get("i_kz_1ph_a"), 1)),
        ("Ток 3ф КЗ, А", _safe_num(res.get("i_kz_3ph_a"), 1)),
        ("Номинал АВ, А", _safe_num(res.get("cb_rating_a"), 0)),
        ("Уставка теплового, А", _safe_num(res.get("cb_thermal_a"), 1)),
        ("Уставка ЭМ, А", _safe_num(res.get("cb_mag_a"), 1)),
        ("Предохранитель ППН, А", _safe_num(res.get("fuse_rating_a"), 0)),
    ])

    # 3. Поправочные коэффициенты
    _add_heading(doc, "3. Поправочные коэффициенты", level=1)
    _add_kv_table(doc, [
        ("k_температура (МЭК B.52.14/B.52.15)", _safe_num(res.get("k_temp"), 3)),
        ("k_группирование (МЭК B.52.17/B.52.18)", _safe_num(res.get("k_group"), 3)),
        ("k_грунт (МЭК B.52.16)", _safe_num(res.get("k_soil"), 3)),
    ])

    # 4. Проверки
    _add_heading(doc, "4. Проверки соответствия нормам", level=1)
    checks = [
        ("Нагрев (I_расч ≤ I_доп)", res.get("check_current")),
        ("Падение напряжения (ΔU ≤ ΔU_доп)", res.get("check_voltage")),
        ("Ток КЗ (срабатывание защиты)", res.get("check_kz")),
    ]
    t = doc.add_table(rows=1 + len(checks), cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Проверка"
    t.rows[0].cells[1].text = "Результат"
    for i, (label, ok) in enumerate(checks, 1):
        t.rows[i].cells[0].text = label
        run_cell = t.rows[i].cells[1].paragraphs[0].add_run(
            "✓ ВЫПОЛНЕНО" if ok else "✗ НЕ ВЫПОЛНЕНО"
        )
        run_cell.bold = True
        run_cell.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a) if ok else RGBColor(0xdc, 0x26, 0x26)

    status = res.get("status", "—")
    p = doc.add_paragraph()
    p.add_run("Итоговый статус: ").bold = True
    r = p.add_run(status)
    r.bold = True
    r.font.color.rgb = (
        RGBColor(0x16, 0xa3, 0x4a) if status == "OK"
        else RGBColor(0xd9, 0x77, 0x06) if status == "WARNING"
        else RGBColor(0xdc, 0x26, 0x26)
    )

    if res.get("hints"):
        _add_heading(doc, "5. Рекомендации", level=1)
        for h in res["hints"]:
            doc.add_paragraph("• " + str(h), style="List Bullet")

    # 6. Методика (формулы)
    methodology = res.get("methodology") or {}
    if methodology:
        _add_heading(doc, "6. Методика расчёта", level=1)
        for key, ru_label in [
            ("i_calc", "Расчётный ток"),
            ("i_allowable", "Допустимый ток"),
            ("delta_u", "Падение напряжения"),
            ("kz", "Ток короткого замыкания"),
            ("protection", "Защита (АВ/предохранитель)"),
        ]:
            block = methodology.get(key)
            if not block:
                continue
            doc.add_heading(ru_label, level=2)
            for fkey in ("formula", "formula_1ph", "formula_3ph"):
                if block.get(fkey):
                    p = doc.add_paragraph()
                    r = p.add_run(block[fkey])
                    r.font.name = "Consolas"
                    r.font.size = Pt(10)
            if block.get("values"):
                rows = list(block["values"].items())
                t2 = doc.add_table(rows=len(rows), cols=2)
                t2.style = "Light List Accent 1"
                for i, (k, v) in enumerate(rows):
                    t2.rows[i].cells[0].text = str(k)
                    t2.rows[i].cells[1].text = "—" if v is None else str(v)
            for rk in ("result_a", "result_pct", "result_1ph_a", "result_3ph_a"):
                if rk in block:
                    p = doc.add_paragraph()
                    p.add_run(f"Результат: {block[rk]}").bold = True
            if block.get("norm"):
                p = doc.add_paragraph()
                run = p.add_run(block["norm"])
                run.italic = True
                run.font.size = Pt(9)

    # 7. Подпись
    doc.add_paragraph()
    doc.add_paragraph()
    sign = doc.add_table(rows=2, cols=2)
    sign.rows[0].cells[0].text = "Расчёт выполнил:"
    sign.rows[0].cells[1].text = "Дата:"
    sign.rows[1].cells[0].text = "_____________________ /                     /"
    sign.rows[1].cells[1].text = time.strftime("%d.%m.%Y")

    p = doc.add_paragraph()
    run = p.add_run("Сформировано в EngHub · cable-calc")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class handler(BaseHTTPRequestHandler):
    def _send_json(self, status: int, payload):
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self._send_json(204, {})

    def do_GET(self):
        self._send_json(200, {
            "ok": True,
            "service": "cable-calc/report",
            "version": "1.0",
            "format": "docx",
        })

    def do_POST(self):
        try:
            if Document is None:
                return self._send_json(500, {
                    "error": f"python-docx not available: {_IMPORT_ERR}"
                })
            length = int(self.headers.get("Content-Length", "0"))
            raw = self.rfile.read(length) if length > 0 else b"{}"
            payload = json.loads(raw.decode("utf-8") or "{}")
            data = _build_doc(payload)
            ts = time.strftime("%Y%m%d_%H%M%S")
            filename = f"cable_calc_{ts}.docx"

            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            self.send_header("Content-Length", str(len(data)))
            self.send_header(
                "Content-Disposition",
                f'attachment; filename="{filename}"',
            )
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Expose-Headers", "Content-Disposition")
            self.end_headers()
            self.wfile.write(data)
        except Exception as e:
            self._send_json(500, {
                "error": str(e),
                "type": type(e).__name__,
                "traceback": traceback.format_exc().splitlines()[-8:],
            })
