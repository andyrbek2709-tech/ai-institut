"""DOCX report generation engine."""

import io
import logging
from typing import Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not available")

from .models import ReportContext, SectionType
from .formula_renderer import FormulaRenderer

logger = logging.getLogger(__name__)


class DocxReportBuilder:
    """Builds production-grade DOCX reports from ReportContext."""

    def __init__(self):
        """Initialize DOCX builder."""
        if not HAS_PYTHON_DOCX:
            raise RuntimeError("python-docx is required for DOCX report generation")

        self.doc = Document()
        self.renderer = FormulaRenderer()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles for engineering reports."""
        # Configure default styles
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        style = self.doc.styles["Heading 1"]
        style.font.name = "Calibri"
        style.font.size = Pt(16)
        style.font.bold = True

        style = self.doc.styles["Heading 2"]
        style.font.name = "Calibri"
        style.font.size = Pt(14)
        style.font.bold = True

    def build_report(self, context: ReportContext) -> bytes:
        """
        Generate complete DOCX from ReportContext.

        Args:
            context: ReportContext with all report data

        Returns:
            DOCX file content as bytes
        """
        try:
            # Build sections in order
            self._add_title_page(context)
            self._add_normative_references(context)
            self._add_assumptions(context)
            self._add_inputs(context)
            self._add_formulas(context)
            self._add_results(context)
            self._add_validation(context)
            self._add_warnings(context)
            self._add_audit_appendix(context)
            self._add_system_info(context)

            # Save to bytes
            output = io.BytesIO()
            self.doc.save(output)
            output.seek(0)
            return output.read()

        except Exception as e:
            logger.error(f"Error building DOCX report: {e}")
            raise

    def _add_title_page(self, context: ReportContext):
        """Add title page section."""
        title = self.doc.add_heading(context.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if context.description:
            p = self.doc.add_paragraph(context.description)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        self.doc.add_paragraph()
        self.doc.add_paragraph(f"Calculation ID: {context.calculation_id}")
        self.doc.add_paragraph(f"Generated: {context.timestamp}")
        self.doc.add_paragraph(f"Template: {context.template_type.value}")

        # Page break
        self.doc.add_page_break()

    def _add_normative_references(self, context: ReportContext):
        """Add normative references section."""
        if not context.normative_references:
            return

        self.doc.add_heading("Normative References", level=2)

        for ref in context.normative_references:
            self.doc.add_paragraph(ref, style="List Bullet")

        self.doc.add_paragraph()

    def _add_assumptions(self, context: ReportContext):
        """Add assumptions section."""
        if not context.assumptions:
            return

        self.doc.add_heading("Assumptions", level=2)

        for assumption in context.assumptions:
            self.doc.add_paragraph(assumption, style="List Bullet")

        self.doc.add_paragraph()

    def _add_inputs(self, context: ReportContext):
        """Add input data section with table."""
        if not context.inputs:
            return

        self.doc.add_heading("Input Data", level=2)

        # Create table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Variable"
        header_cells[1].text = "Value"
        header_cells[2].text = "Unit"
        header_cells[3].text = "Description"

        # Data rows
        for var_name, var in context.inputs.items():
            row_cells = table.add_row().cells
            row_cells[0].text = var.label
            row_cells[1].text = str(var.value)
            row_cells[2].text = var.unit or "—"
            row_cells[3].text = var.description or "—"

        self.doc.add_paragraph()

    def _add_formulas(self, context: ReportContext):
        """Add formulas and calculations section."""
        if not context.formulas:
            return

        self.doc.add_heading("Formulas & Calculations", level=2)

        for idx, formula in enumerate(context.formulas, 1):
            self.doc.add_heading(f"{idx}. {formula.name}", level=3)

            if formula.description:
                self.doc.add_paragraph(formula.description)

            # Formula
            self.doc.add_paragraph(f"Formula: {formula.expression}")

            # LaTeX version (if available)
            if formula.latex_formula and formula.latex_formula != formula.expression:
                self.doc.add_paragraph(f"LaTeX: {formula.latex_formula}")

            # Variable definitions
            if formula.variable_definitions:
                self.doc.add_paragraph("Where:")
                for line in formula.variable_definitions.split("\n"):
                    if line.strip():
                        self.doc.add_paragraph(line, style="List Bullet")

            # Calculation steps
            if formula.calculation_steps:
                self.doc.add_paragraph("Calculation:")
                for step in formula.calculation_steps:
                    p = self.doc.add_paragraph(step)
                    p.paragraph_format.left_indent = Inches(0.25)

            # Result
            if formula.output_unit:
                self.doc.add_paragraph(
                    f"Result: {formula.output_value} {formula.output_unit}"
                )
            else:
                self.doc.add_paragraph(f"Result: {formula.output_value}")

            # Source reference
            if formula.source_reference:
                p = self.doc.add_paragraph(f"Reference: {formula.source_reference}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.italic = True

            self.doc.add_paragraph()

    def _add_results(self, context: ReportContext):
        """Add results summary section."""
        if not context.results:
            return

        self.doc.add_heading("Results", level=2)

        # Create results table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Output Variable"
        header_cells[1].text = "Value"
        header_cells[2].text = "Unit"

        # Data rows
        for var_name, result in context.results.items():
            row_cells = table.add_row().cells
            row_cells[0].text = result.label
            row_cells[1].text = str(result.value)
            row_cells[2].text = result.unit or "—"

            # Highlight critical results
            if result.is_critical:
                for cell in row_cells:
                    self._highlight_cell(cell, "FFFF00")  # Yellow

        self.doc.add_paragraph()

    def _add_validation(self, context: ReportContext):
        """Add validation results section."""
        if not context.validation_results:
            return

        self.doc.add_heading("Validation Results", level=2)

        # Summary
        passed = sum(1 for v in context.validation_results if v.status == "passed")
        failed = sum(1 for v in context.validation_results if v.status == "failed")

        summary = f"Total: {len(context.validation_results)} | Passed: {passed} | Failed: {failed}"
        p = self.doc.add_paragraph(summary)
        p.runs[0].font.bold = True

        # Validation table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Rule"
        header_cells[1].text = "Status"
        header_cells[2].text = "Severity"
        header_cells[3].text = "Message"

        # Data rows
        for validation in context.validation_results:
            row_cells = table.add_row().cells
            row_cells[0].text = validation.rule_name
            row_cells[1].text = validation.status.upper()
            row_cells[2].text = validation.severity.upper()
            row_cells[3].text = validation.message

            # Color code by status
            if validation.status == "passed":
                color = "90EE90"  # Light green
            elif validation.status == "failed":
                color = "FFB6C6"  # Light red
            else:
                color = "FFFFE0"  # Light yellow

            for cell in row_cells:
                self._highlight_cell(cell, color)

        self.doc.add_paragraph()

    def _add_warnings(self, context: ReportContext):
        """Add engineering warnings section."""
        if not context.warnings:
            return

        self.doc.add_heading("Engineering Warnings", level=2)

        for warning in context.warnings:
            p = self.doc.add_paragraph(warning, style="List Bullet")
            p.runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange

        self.doc.add_paragraph()

    def _add_audit_appendix(self, context: ReportContext):
        """Add audit appendix with execution traces."""
        if (
            not context.execution_traces
            and not context.failure_analysis
            and not context.audit_trail
        ):
            return

        self.doc.add_page_break()
        self.doc.add_heading("Appendix A: Calculation Audit Trail", level=2)

        # Failure analysis
        if context.failure_analysis:
            self.doc.add_heading("Failure Analysis", level=3)
            self.doc.add_paragraph(context.failure_analysis.summary_text)

            for failure in context.failure_analysis.failures:
                self.doc.add_paragraph(
                    f"• {failure.get('rule', 'Unknown rule')}: {failure.get('message', '')}"
                )

        # Execution traces
        if context.execution_traces:
            self.doc.add_heading("Execution Traces", level=3)

            for trace in context.execution_traces:
                self.doc.add_paragraph(f"Formula: {trace.formula_id}")
                self.doc.add_paragraph(f"Expression: {trace.expression}")
                self.doc.add_paragraph(f"Output: {trace.output} {trace.unit or ''}")
                self.doc.add_paragraph(f"Duration: {trace.duration_ms:.2f} ms")
                self.doc.add_paragraph()

        self.doc.add_paragraph()

    def _add_system_info(self, context: ReportContext):
        """Add system information appendix."""
        self.doc.add_page_break()
        self.doc.add_heading("Appendix B: System Information", level=2)

        self.doc.add_paragraph(
            f"Calculation Engine Version: {context.calculation_engine_version}"
        )
        if context.calculation_time_ms:
            self.doc.add_paragraph(
                f"Calculation Time: {context.calculation_time_ms:.2f} ms"
            )
        self.doc.add_paragraph(f"Generated: {context.timestamp}")
        self.doc.add_paragraph(f"Calculation ID: {context.calculation_id}")

    @staticmethod
    def _highlight_cell(cell, color: str):
        """Highlight a table cell with specified color."""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
